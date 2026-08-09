"""Reaction-scheme renderer.

Composes ONE organic reaction into a single sharp **vector SVG** (+ PNG fallback) by
reusing the per-molecule engine in structure_svg / structure_draw:

    reactant(s)  + ... + reactant   --[reagent above / conditions below]-->   product(s)

Species are passed as strings:
  * a leading "$"  -> LITERAL text token (inorganics / ions / radicals / coefficients).
                       Digits auto-subscript after a letter / ")" / "]"  (H2O -> H₂O,
                       Br2 -> Br₂, K2Cr2O7 -> K₂Cr₂O₇); a LEADING digit stays a full-size
                       coefficient (2[O], 3OH). Use "^" for a charge / superscript
                       ("3OH^-", "2[Ag(NH3)2]^+", "Na^+"); "_" forces a subscript.
                       "½" and "•" are written literally.  e.g. "$NaOH(aq)", "$2[O]",
                       "$•CH3", "$2[Ag(NH3)2]^+", "$½H2"
  * anything else  -> SMILES, drawn by structure_draw  (e.g. "CCO", "CC(=O)[O-]", "c1ccccc1").
                       Charges in SMILES must be ASCII: "[O-]", "[N+]".

Arrows:  arrow="->"  straight reaction arrow (filled head)
         arrow="<=>" equilibrium (two opposed harpoons)

Public API:
    render_reaction_svg(reactants, products, reagent="", conditions="", arrow="->") -> (svg, W, H)
    render_reaction_png(reactants, products, ...) -> PIL.Image
    add_reaction(doc, run, reactants, products, ..., width_in=4.5)   # SVG (vector) + PNG fallback
    build_reaction_doc(items, out_path, title=None, width_in=4.5)    # one reaction per row, captioned
"""
import os, sys, re, math, tempfile
sys.path.insert(0, os.path.dirname(__file__))
import structure_draw as J
import structure_svg as S
from structure_draw import U, STROKE, DBL, FONT_PATH, FONT_NAME, _BASE_H, _txt_w
from PIL import Image, ImageDraw, ImageFont

FS = S.FS            # 34: main glyph, matches molecule labels
RXN_FS = 26          # reagent / conditions text above & below the arrow
LINE_H = int(RXN_FS * 1.38)
PAD = 14             # outer canvas padding
GAP = 10             # horizontal gap between adjacent blocks
ARROW_MIN = 150      # minimum arrow shaft length (px)
SUB_RATIO = 0.68     # subscript / superscript size relative to its run
SUB_DY = 0.24        # subscript baseline drop (× run size)
SUP_DY = -0.42       # superscript baseline rise (× run size)
CENTRE = 0.35        # baseline offset to vertically centre a glyph on a point

# ----------------------------------------------------------------------------- fonts
_font_cache = {}
def _font(sz):
    if sz not in _font_cache:
        _font_cache[sz] = ImageFont.truetype(FONT_PATH, sz)
    return _font_cache[sz]

def _esc(s):
    return s.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')

def _text_w(text, sz):
    return _txt_w(text, _font(sz))

# ----------------------------------------------------------------------------- formula text
_FRAC_GLYPHS = {'½': '1/2', '⅓': '1/3', '⅔': '2/3', '¼': '1/4', '¾': '3/4',
                '⅕': '1/5', '⅖': '2/5', '⅗': '3/5', '⅘': '4/5', '⅙': '1/6',
                '⅛': '1/8', '⅜': '3/8', '⅝': '5/8', '⅞': '7/8'}

def _preprocess_fracs(text):
    for g, r in _FRAC_GLYPHS.items():
        text = text.replace(g, r)
    return text

def _formula_runs(text):
    """Split a literal token into (substring|tuple, kind) runs.
    kind in {'n','sub','sup','frac'}. A numeric d/d becomes a built-up fraction."""
    text = _preprocess_fracs(text)
    runs = []
    i, n, prev = 0, len(text), ''
    while i < n:
        c = text[i]
        if c == '^':                                   # explicit superscript
            i += 1
            if i < n and text[i] == '{':
                j = text.index('}', i); runs.append((text[i+1:j], 'sup')); i = j + 1
            else:
                m = re.match(r'[0-9]+[+\-−]?|[+\-−]|.', text[i:]); s = m.group(0)
                runs.append((s, 'sup')); i += len(s)
            prev = ''; continue
        if c == '_':                                   # explicit subscript
            i += 1
            if i < n and text[i] == '{':
                j = text.index('}', i); runs.append((text[i+1:j], 'sub')); i = j + 1
            else:
                m = re.match(r'[0-9]+|.', text[i:]); s = m.group(0)
                runs.append((s, 'sub')); i += len(s)
            prev = ''; continue
        if c in '0123456789':                              # ASCII digit only (Unicode ₂₃ fall through)
            fm = re.match(r'([0-9]+)/([0-9]+)', text[i:])   # built-up fraction
            if fm:
                runs.append(((fm.group(1), fm.group(2)), 'frac'))
                i += len(fm.group(0)); prev = ''; continue
            m = re.match(r'[0-9]+', text[i:]); s = m.group(0)
            kind = 'sub' if (prev and (prev.isalpha() or prev in ')]')) else 'n'
            runs.append((s, kind)); i += len(s); prev = s[-1]; continue
        runs.append((c, 'n'))
        if not c.isspace(): prev = c
        i += 1
    merged = []                                        # coalesce adjacent normal runs
    for s, k in runs:
        if merged and merged[-1][1] == 'n' and k == 'n':
            merged[-1] = (merged[-1][0] + s, 'n')
        else:
            merged.append((s, k))
    return merged

def _runs_w(runs, main_sz):
    sub = max(11, int(main_sz * SUB_RATIO))
    w = 0
    for s, k in runs:
        if k == 'frac':
            num, den = s
            w += max(_text_w(num, sub), _text_w(den, sub)) + 6
        else:
            w += _text_w(s, main_sz if k == 'n' else sub)
    return w

def _runs_svg_pil(runs, cx, lc, main_sz):
    """Render formula runs centred horizontally on cx, vertically on lc -> (svg, pil_fn)."""
    sub = max(11, int(main_sz * SUB_RATIO))
    base_n = lc + main_sz * CENTRE
    x = cx - _runs_w(runs, main_sz) / 2.0
    svg, ops = [], []

    def _txt(s, xx, base, sz):
        svg.append(f'<text x="{xx:.1f}" y="{base:.1f}" font-family="{FONT_NAME}" '
                   f'font-size="{sz}" text-anchor="start">{_esc(s)}</text>')
        ops.append(lambda img, dr, ox, oy, s=s, xx=xx, base=base, sz=sz:
                   dr.text((ox + xx, oy + base), s, fill=(0, 0, 0), font=_font(sz), anchor='ls'))

    for s, k in runs:
        if k == 'frac':
            num, den = s
            fw = max(_text_w(num, sub), _text_w(den, sub)) + 6
            _txt(num, x + (fw - _text_w(num, sub)) / 2.0, lc - 3, sub)
            _txt(den, x + (fw - _text_w(den, sub)) / 2.0, lc + sub * 0.92, sub)
            bx0, bx1, by = x + 1, x + fw - 1, lc
            svg.append(f'<line x1="{bx0:.1f}" y1="{by:.1f}" x2="{bx1:.1f}" y2="{by:.1f}" '
                       f'stroke="black" stroke-width="2"/>')
            ops.append(lambda img, dr, ox, oy, bx0=bx0, bx1=bx1, by=by:
                       dr.line([(ox + bx0, oy + by), (ox + bx1, oy + by)], fill=(0, 0, 0), width=2))
            x += fw
        else:
            sz = main_sz if k == 'n' else sub
            base = base_n + (main_sz * SUB_DY if k == 'sub' else main_sz * SUP_DY if k == 'sup' else 0)
            _txt(s, x, base, sz)
            x += _text_w(s, sz)

    def pil(img, dr, ox, oy):
        for op in ops:
            op(img, dr, ox, oy)
    return "".join(svg), pil

def _stack_lines(lines, cx, top, sz):
    """Stacked centred text lines from `top` downward -> (svg, pil_fn)."""
    svgs, pilfns = [], []
    for i, ln in enumerate(lines):
        s, f = _runs_svg_pil(_formula_runs(ln), cx, top + (i + 0.5) * LINE_H, sz)
        svgs.append(s); pilfns.append(f)
    def pil(img, dr, ox, oy):
        for f in pilfns: f(img, dr, ox, oy)
    return "".join(svgs), pil

# A "block" = dict(w, h, cy, svg, pil)
#   cy : the y within the block that lands on the common reaction axis
#   pil: callable(img, draw, ox, oy) -> paints this block onto a PIL canvas at (ox,oy)

def _centered_text_block(text, sz=FS):
    runs = _formula_runs(text)
    w = int(_runs_w(runs, sz)) + 12
    h = int(_BASE_H + 22)
    cy = h / 2.0
    svg, pilfn = _runs_svg_pil(runs, w / 2.0, cy, sz)
    return dict(w=w, h=h, cy=cy, svg=svg, pil=pilfn)

def _molecule_block(smiles):
    mol = J.validate_input(smiles)
    inner, w, h, px = S._render_svg_inner_mol(mol)
    cy = px(0.0, 0.0)[1]                       # backbone row (coord y=0) -> align here
    pil_img = J._render_mol(mol)               # same validated molecule and geometry as the SVG (white bg)
    def pil(img, dr, ox, oy):
        img.paste(pil_img, (int(round(ox)), int(round(oy))))
    return dict(w=w, h=h, cy=cy, svg=inner, pil=pil)

def _plus_block():
    return _centered_text_block("+", FS)

def _arrow_block(reagent="", conditions="", arrow="->", arrow_len=ARROW_MIN):
    rg = reagent.split("\n") if reagent else []
    cn = conditions.split("\n") if conditions else []
    rg_w = max([_runs_w(_formula_runs(l), RXN_FS) for l in rg], default=0)
    cn_w = max([_runs_w(_formula_runs(l), RXN_FS) for l in cn], default=0)
    width = int(max(arrow_len, rg_w + 28, cn_w + 28))
    rg_h, cn_h = len(rg) * LINE_H, len(cn) * LINE_H
    GAPV = 8
    cy = rg_h + GAPV                            # shaft y from block top
    h = max(int(cy + GAPV + cn_h), int(cy + GAPV) + 2)
    cx = width / 2.0
    x0, x1 = 3.0, width - 3.0

    elems = []
    rg_pil = cn_pil = None
    if rg:
        s, rg_pil = _stack_lines(rg, cx, 0, RXN_FS); elems.append(s)
    if cn:
        s, cn_pil = _stack_lines(cn, cx, cy + GAPV, RXN_FS); elems.append(s)

    def _arrow_svg():
        out = []
        if arrow == "<=>":
            yt, yb = cy - 3.5, cy + 3.5
            out.append(f'<line x1="{x0:.1f}" y1="{yt:.1f}" x2="{x1:.1f}" y2="{yt:.1f}" stroke="black" stroke-width="{STROKE}" stroke-linecap="round"/>')
            out.append(f'<line x1="{x1:.1f}" y1="{yt:.1f}" x2="{x1-11:.1f}" y2="{yt-5:.1f}" stroke="black" stroke-width="{STROKE}" stroke-linecap="round"/>')
            out.append(f'<line x1="{x1:.1f}" y1="{yb:.1f}" x2="{x0:.1f}" y2="{yb:.1f}" stroke="black" stroke-width="{STROKE}" stroke-linecap="round"/>')
            out.append(f'<line x1="{x0:.1f}" y1="{yb:.1f}" x2="{x0+11:.1f}" y2="{yb+5:.1f}" stroke="black" stroke-width="{STROKE}" stroke-linecap="round"/>')
        else:
            out.append(f'<line x1="{x0:.1f}" y1="{cy:.1f}" x2="{x1:.1f}" y2="{cy:.1f}" stroke="black" stroke-width="{STROKE}" stroke-linecap="round"/>')
            out.append(f'<path d="M {x1:.1f} {cy:.1f} L {x1-12:.1f} {cy-5.5:.1f} L {x1-12:.1f} {cy+5.5:.1f} Z" fill="black"/>')
        return "".join(out)
    elems.append(_arrow_svg())
    svg = "".join(elems)

    def pil(img, dr, ox, oy):
        if rg_pil: rg_pil(img, dr, ox, oy)
        if cn_pil: cn_pil(img, dr, ox, oy)
        if arrow == "<=>":
            yt, yb = cy - 3.5, cy + 3.5
            dr.line([(ox+x0, oy+yt), (ox+x1, oy+yt)], fill=(0,0,0), width=STROKE)
            dr.line([(ox+x1, oy+yt), (ox+x1-11, oy+yt-5)], fill=(0,0,0), width=STROKE)
            dr.line([(ox+x1, oy+yb), (ox+x0, oy+yb)], fill=(0,0,0), width=STROKE)
            dr.line([(ox+x0, oy+yb), (ox+x0+11, oy+yb+5)], fill=(0,0,0), width=STROKE)
        else:
            dr.line([(ox+x0, oy+cy), (ox+x1, oy+cy)], fill=(0,0,0), width=STROKE)
            dr.polygon([(ox+x1, oy+cy), (ox+x1-12, oy+cy-5.5), (ox+x1-12, oy+cy+5.5)], fill=(0,0,0))

    return dict(w=width, h=h, cy=cy, svg=svg, pil=pil)

def _species_block(token):
    token = token.strip()
    if token.startswith("$"):
        return _centered_text_block(token[1:].strip(), FS)
    return _molecule_block(token)

def _side_blocks(species):
    blocks = []
    for k, sp in enumerate(species):
        if k:
            blocks.append(_plus_block())
        blocks.append(_species_block(sp))
    return blocks

def _compose(reactants, products, reagent, conditions, arrow, arrow_len):
    blocks = _side_blocks(reactants)
    blocks.append(_arrow_block(reagent, conditions, arrow, arrow_len))
    blocks += _side_blocks(products)
    CY = max(b["cy"] for b in blocks)
    x = PAD
    placed = []
    for b in blocks:
        ty = PAD + (CY - b["cy"])
        placed.append((b, x, ty))
        x += b["w"] + GAP
    W = int(x - GAP + PAD)
    H = int(max(ty + b["h"] for b, _, ty in placed) + PAD)
    return placed, W, H


def _placed_svg(placed, W, H):
    body = [f'<rect width="{W}" height="{H}" fill="white"/>']
    for b, tx, ty in placed:
        body.append(f'<g transform="translate({tx:.1f},{ty:.1f})">{b["svg"]}</g>')
    return (f'<svg xmlns="http://www.w3.org/2000/svg" width="{W}" height="{H}" '
            f'viewBox="0 0 {W} {H}">{"".join(body)}</svg>')


def _placed_png(placed, W, H):
    img = Image.new("RGB", (W, H), "white")
    dr = ImageDraw.Draw(img)
    for b, tx, ty in placed:
        b["pil"](img, dr, tx, ty)
    return img


def render_reaction_svg(reactants, products, reagent="", conditions="",
                        arrow="->", arrow_len=ARROW_MIN):
    """reactants, products: list[str] of species ($literal or SMILES). -> (svg, W, H)."""
    placed, W, H = _compose(reactants, products, reagent, conditions, arrow, arrow_len)
    return _placed_svg(placed, W, H), W, H

def render_reaction_png(reactants, products, reagent="", conditions="",
                        arrow="->", arrow_len=ARROW_MIN):
    placed, W, H = _compose(reactants, products, reagent, conditions, arrow, arrow_len)
    return _placed_png(placed, W, H)

# ----------------------------------------------------------------------------- docx
_SVG_EXT_URI = "{96DAC541-7B7A-43D3-8B79-37D633B846F1}"
_counter = [0]

def add_reaction(doc, run, reactants, products, reagent="", conditions="",
                 arrow="->", width_in=4.5, arrow_len=ARROW_MIN, png_path=None):
    """Render the scheme into `run` as SVG (vector) + PNG fallback (MS Office SVG ext)."""
    from docx.shared import Inches
    from docx.opc.packuri import PackURI
    from docx.opc.part import Part
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    from docx.oxml import parse_xml
    placed, W, H = _compose(reactants, products, reagent, conditions, arrow, arrow_len)
    svg = _placed_svg(placed, W, H)
    if png_path is None:
        png_path = os.path.join(tempfile.gettempdir(), f"_rxn_{_counter[0]}.png")
    _placed_png(placed, W, H).save(png_path)
    run.add_picture(png_path, width=Inches(width_in))
    part = Part(PackURI(f'/word/media/rxn{_counter[0]}.svg'),
                'image/svg+xml', svg.encode('utf-8'), doc.part.package)
    _counter[0] += 1
    rid = doc.part.relate_to(part, RT.IMAGE)
    blip = run._r.xpath('.//a:blip')[-1]
    blip.append(parse_xml(
        '<a:extLst xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">'
        f'<a:ext uri="{_SVG_EXT_URI}">'
        '<asvg:svgBlip xmlns:asvg="http://schemas.microsoft.com/office/drawing/2016/SVG/main" '
        'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" '
        f'r:embed="{rid}"/></a:ext></a:extLst>'))

def build_reaction_doc(items, out_path, title=None, width_in=4.5):
    """items: list of dicts {reactants, products, [reagent, conditions, arrow, caption]}.
    One reaction per row (full width) with an optional caption beneath."""
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    doc = Document()
    if title:
        doc.add_heading(title, 0)
    for it in items:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_reaction(doc, p.add_run(),
                     it["reactants"], it["products"],
                     it.get("reagent", ""), it.get("conditions", ""),
                     it.get("arrow", "->"), width_in=width_in,
                     arrow_len=it.get("arrow_len", ARROW_MIN))
        cap = it.get("caption")
        if cap:
            pc = doc.add_paragraph(); pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = pc.add_run(cap); r.font.size = Pt(10); r.italic = True
        doc.add_paragraph()
    doc.save(out_path)
    return out_path
