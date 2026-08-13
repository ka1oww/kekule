"""Regression tests for Kekule's public SMILES validation boundary."""

import hashlib
import json
import math
import os
import sys
import unittest
import xml.etree.ElementTree as ET
from unittest import mock

from PIL import Image, ImageChops

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, ROOT)

import reaction_render as reaction
import embed_docx
import structure_draw as draw
import structure_svg as svg


class ValidationBoundaryTests(unittest.TestCase):
    @staticmethod
    def _reference_sides(smiles):
        mol = draw.validate_input(smiles, "displayed")
        double = next(
            bond for bond in mol.GetBonds()
            if bond.GetBondType() == draw.Chem.BondType.DOUBLE and not bond.IsInRing()
        )
        c1, c2 = double.GetBeginAtomIdx(), double.GetEndAtomIdx()
        ref1, ref2 = double.GetStereoAtoms()
        atoms, _, _, _ = draw.layout(smiles, "displayed")
        ax = atoms[c2][1] - atoms[c1][1]
        ay = atoms[c2][2] - atoms[c1][2]

        def side(centre, reference):
            vx = atoms[reference][1] - atoms[centre][1]
            vy = atoms[reference][2] - atoms[centre][2]
            return 1 if ax * vy - ay * vx > 0 else -1

        return side(c1, ref1), side(c2, ref2)

    @staticmethod
    def _angle(atoms, centre, first, second):
        ax = atoms[first][1] - atoms[centre][1]
        ay = atoms[first][2] - atoms[centre][2]
        bx = atoms[second][1] - atoms[centre][1]
        by = atoms[second][2] - atoms[centre][2]
        cosine = (ax * bx + ay * by) / (math.hypot(ax, ay) * math.hypot(bx, by))
        return math.degrees(math.acos(max(-1.0, min(1.0, cosine))))

    def test_valid_core_input_returns_a_reusable_molecule(self):
        mol = draw.validate_input("CC(=O)O")
        self.assertEqual(mol.GetNumAtoms(), 4)
        self.assertEqual(mol.GetNumBonds(), 3)
        for form in ("structural", "displayed", "skeletal"):
            atoms, bonds, _, _ = draw._layout_mol(mol, form)
            self.assertTrue(atoms, form)
            self.assertTrue(bonds, form)

    def test_chemical_name_is_not_treated_as_smiles(self):
        with self.assertRaisesRegex(draw.InvalidSmilesError, "not chemical names"):
            draw.layout("ethanoic acid")

    def test_malformed_smiles_has_a_typed_human_readable_error(self):
        with self.assertRaisesRegex(draw.InvalidSmilesError, "Invalid SMILES"):
            draw.render("not-smiles")

    def test_empty_or_non_string_input_is_invalid(self):
        for value in ("", "   ", None, 42):
            with self.subTest(value=value):
                with self.assertRaises(draw.InvalidSmilesError):
                    draw.validate_input(value)

    def test_disconnected_stereo_request_is_rejected(self):
        with self.assertRaisesRegex(draw.DisconnectedStructureError, "stereo renderer requires one"):
            draw.render("CC(O)C.[Na+]", form="stereo")

    def test_unsupported_fragment_error_names_the_fragment(self):
        smiles = "[Na+].c1ccc2ccccc2c1"
        with self.assertRaisesRegex(draw.UnsupportedTopologyError, r"Fragment 2 'c1ccc2ccccc2c1'"):
            draw.render(smiles)

    def test_radical_is_rejected_until_the_electron_dot_is_supported(self):
        with self.assertRaisesRegex(draw.UnsupportedRadicalError, "unpaired-electron dots"):
            svg.render_svg("[CH3]")

    def test_every_unsupported_multi_ring_topology_is_rejected(self):
        cases = {
            "naphthalene": "c1ccc2ccccc2c1",
            "biphenyl": "c1ccccc1-c2ccccc2",
            "decalin": "C1CCC2CCCCC2C1",
            "norbornane": "C1CC2CCC1C2",
            "spiro": "C1CCC2(CC1)CCCC2",
        }
        for name, smiles in cases.items():
            with self.subTest(name=name):
                with self.assertRaisesRegex(draw.UnsupportedTopologyError, "at most one simple ring"):
                    draw.layout(smiles)

    def test_multi_centre_stereo_pair_is_rejected(self):
        smiles = "O=C(O)[C@H](O)[C@H](O)C(=O)O"
        with self.assertRaisesRegex(draw.UnsupportedStereochemistryError, "exactly one"):
            draw.render_stereo_pair(smiles)

    def test_supported_stereo_and_geometric_pairs_still_render(self):
        stereo = draw.render_stereo_pair("CC(O)C(=O)O")
        geometric = draw.render_geometric_pair("CC=CC")
        images = [stereo] + [image for _, image in geometric]
        for image in images:
            with self.subTest(size=image.size):
                self.assertGreater(image.width, 0)
                self.assertGreater(image.height, 0)

    def test_directional_alkene_stereo_changes_displayed_geometry(self):
        self.assertEqual(self._reference_sides("C/C=C\\C")[0], self._reference_sides("C/C=C\\C")[1])
        self.assertNotEqual(self._reference_sides("C/C=C/C")[0], self._reference_sides("C/C=C/C")[1])

    def test_directional_ring_alkene_stereo_changes_displayed_geometry(self):
        e_sides = self._reference_sides("c1ccccc1/C=C/Cl")
        z_sides = self._reference_sides("c1ccccc1/C=C\\Cl")
        self.assertNotEqual(e_sides[0], e_sides[1])
        self.assertEqual(z_sides[0], z_sides[1])
        self.assertNotEqual(e_sides, z_sides)

    def test_skeletal_form_keeps_formal_charge_on_carbon(self):
        for smiles in ("[CH2+]CBr", "C[CH+]CC", "[CH+]1CCCCC1"):
            atoms, _, _, _ = draw.layout(smiles, "skeletal")
            self.assertTrue(any(label.startswith("C^") for label, _, _ in atoms.values()))

    def test_displayed_ring_branches_preserve_public_topology(self):
        cases = {
            "amide": "CC(=O)Nc1ccc(O)cc1",
            "ester": "CCOC(=O)c1ccccc1",
            "aldehyde": "O=Cc1ccccc1",
            "methoxy": "COc1ccccc1",
            "benzyl": "OCc1ccccc1",
            "branched alkyl": "CC(C)(C)c1ccccc1",
        }
        for name, smiles in cases.items():
            with self.subTest(name=name):
                mol = draw.validate_input(smiles, "displayed")
                mol_h = draw.Chem.AddHs(mol)
                ring = set(mol.GetRingInfo().AtomRings()[0])
                expected_atoms = set(ring)
                for atom in mol_h.GetAtoms():
                    idx = atom.GetIdx()
                    if atom.GetAtomicNum() > 1 and idx not in ring:
                        expected_atoms.add(idx)
                    elif atom.GetAtomicNum() == 1 and atom.GetNeighbors()[0].GetIdx() not in ring:
                        expected_atoms.add(idx)
                expected_bonds = {
                    (min(bond.GetBeginAtomIdx(), bond.GetEndAtomIdx()),
                     max(bond.GetBeginAtomIdx(), bond.GetEndAtomIdx()),
                     draw._BONDORD.get(bond.GetBondType(), 1))
                    for bond in mol_h.GetBonds()
                    if bond.GetBeginAtomIdx() in expected_atoms and bond.GetEndAtomIdx() in expected_atoms
                }
                atoms, bonds, _, _ = draw.layout(smiles, "displayed")
                actual_bonds = {(min(first, second), max(first, second), order)
                                for first, second, order in bonds}
                self.assertEqual(set(atoms), expected_atoms)
                self.assertEqual(actual_bonds, expected_bonds)

    def test_displayed_ring_branches_use_bent_and_trigonal_slots(self):
        methoxy = "COc1ccccc1"
        methoxy_mol = draw.validate_input(methoxy, "displayed")
        methoxy_h = draw.Chem.AddHs(methoxy_mol)
        methoxy_atoms, _, _, _ = draw.layout(methoxy, "displayed")
        oxygen = next(atom.GetIdx() for atom in methoxy_h.GetAtoms()
                      if atom.GetAtomicNum() == 8 and atom.GetDegree() == 2)
        oxygen_neighbours = [atom.GetIdx() for atom in methoxy_h.GetAtomWithIdx(oxygen).GetNeighbors()]
        self.assertAlmostEqual(
            self._angle(methoxy_atoms, oxygen, *oxygen_neighbours), 120.0, delta=1.0
        )
        methyl = next(idx for idx in oxygen_neighbours
                      if methoxy_h.GetAtomWithIdx(idx).GetAtomicNum() == 6
                      and idx not in methoxy_mol.GetRingInfo().AtomRings()[0])
        parent = oxygen
        hydrogens = [atom.GetIdx() for atom in methoxy_h.GetAtomWithIdx(methyl).GetNeighbors()
                     if atom.GetAtomicNum() == 1]
        self.assertTrue(all(self._angle(methoxy_atoms, methyl, parent, hydrogen) > 45
                            for hydrogen in hydrogens))

        amide = "CC(=O)Nc1ccccc1"
        amide_mol = draw.validate_input(amide, "displayed")
        amide_h = draw.Chem.AddHs(amide_mol)
        amide_atoms, _, _, _ = draw.layout(amide, "displayed")
        nitrogen = next(atom.GetIdx() for atom in amide_h.GetAtoms() if atom.GetAtomicNum() == 7)
        nitrogen_neighbours = [atom.GetIdx() for atom in amide_h.GetAtomWithIdx(nitrogen).GetNeighbors()]
        angles = [self._angle(amide_atoms, nitrogen, nitrogen_neighbours[i], nitrogen_neighbours[j])
                  for i in range(len(nitrogen_neighbours))
                  for j in range(i + 1, len(nitrogen_neighbours))]
        self.assertTrue(all(abs(angle - 120.0) <= 1.0 for angle in angles))

        aldehyde = "O=Cc1ccccc1"
        aldehyde_mol = draw.validate_input(aldehyde, "displayed")
        aldehyde_h = draw.Chem.AddHs(aldehyde_mol)
        aldehyde_atoms, _, _, _ = draw.layout(aldehyde, "displayed")
        carbonyl = next(atom.GetIdx() for atom in aldehyde_h.GetAtoms()
                        if atom.GetAtomicNum() == 6
                        and any(bond.GetBondType() == draw.Chem.BondType.DOUBLE for bond in atom.GetBonds()))
        oxygen = next(atom.GetIdx() for atom in aldehyde_h.GetAtomWithIdx(carbonyl).GetNeighbors()
                      if atom.GetAtomicNum() == 8)
        hydrogen = next(atom.GetIdx() for atom in aldehyde_h.GetAtomWithIdx(carbonyl).GetNeighbors()
                        if atom.GetAtomicNum() == 1)
        self.assertAlmostEqual(
            self._angle(aldehyde_atoms, carbonyl, oxygen, hydrogen), 120.0, delta=1.0
        )

    def test_chlorine_lowercase_l_is_italic_in_svg(self):
        markup, _, _ = svg.render_svg("Clc1ccccc1")
        self.assertIn('font-style="italic">l</tspan>', markup)

        reaction_markup, _, _ = reaction.render_reaction_svg(
            ["$HCl"], ["$Cl2"], reagent="AlCl3", conditions="HCl"
        )
        self.assertEqual(reaction_markup.count('font-style="italic">l</tspan>'), 4)
        with mock.patch.object(draw, "_draw_inline", wraps=draw._draw_inline) as inline:
            image = reaction.render_reaction_png(
                ["$HCl"], ["$Cl2"], reagent="AlCl3", conditions="HCl"
            )
        self.assertGreater(image.width, 0)
        self.assertTrue(any("Cl" in call.args[3] for call in inline.call_args_list))

    def test_left_facing_condensed_labels_reverse_the_non_anchor_atoms(self):
        self.assertEqual(draw._suffix_for_side("C", "HO", "left"), "OH")
        self.assertEqual(draw._suffix_for_side("C", "OOH", "left"), "HOO")
        self.assertEqual(draw._suffix_for_side("N", "H2", "left"), "H2")

    def test_non_stereogenic_pair_requests_have_typed_errors(self):
        with self.assertRaises(draw.UnsupportedStereochemistryError):
            draw.render_stereo_pair("CC")
        with self.assertRaisesRegex(draw.UnsupportedStereochemistryError, "two different groups"):
            draw.render_geometric_pair("C=C")

    def test_stereo_svg_has_a_typed_unsupported_error(self):
        for call in (svg.render_svg, svg.render_svg_inner):
            with self.subTest(call=call):
                with self.assertRaisesRegex(draw.UnsupportedStereochemistryError, "SVG backend"):
                    call("CC(O)C(=O)O", form="stereo")

    def test_supported_simple_rings_still_render(self):
        for smiles in ("c1ccccc1", "C1CCCCC1", "C1=CCCCC1"):
            for form in ("structural", "displayed", "skeletal"):
                with self.subTest(smiles=smiles, form=form):
                    markup, width, height = svg.render_svg(smiles, form=form)
                    self.assertTrue(markup.startswith("<svg"))
                    self.assertGreater(width, 0)
                    self.assertGreater(height, 0)


class MultiFragmentRenderingTests(unittest.TestCase):
    @staticmethod
    def _components(atoms, bonds):
        adjacency = {atom_id: set() for atom_id in atoms}
        for first, second, _ in bonds:
            adjacency[first].add(second)
            adjacency[second].add(first)
        components = []
        unseen = set(atoms)
        while unseen:
            pending = [next(iter(unseen))]
            component = set()
            while pending:
                atom_id = pending.pop()
                if atom_id in component:
                    continue
                component.add(atom_id)
                pending.extend(adjacency[atom_id] - component)
            unseen -= component
            components.append(component)
        return components

    def test_sodium_ethanoate_fragments_have_a_clear_non_overlapping_gap(self):
        for form in ("displayed", "structural", "skeletal"):
            with self.subTest(form=form):
                atoms, bonds, circles, reversible = draw.layout("CC(=O)[O-].[Na+]", form)
                components = self._components(atoms, bonds)
                self.assertEqual(len(components), 2)
                bounds = []
                for component in components:
                    component_atoms = {i: atoms[i] for i in component}
                    component_bonds = [bond for bond in bonds if bond[0] in component]
                    bounds.append(draw._layout_visual_bounds(
                        component_atoms, component_bonds, [], reversible & component
                    ))
                bounds.sort(key=lambda bound: bound[0])
                self.assertGreaterEqual(bounds[1][0] - bounds[0][1], draw.FRAGMENT_GAP - 1e-9)

    def test_monatomic_ion_labels_and_charges_are_preserved(self):
        structural = draw.layout("[NH4+].[Cl-]", "structural")
        self.assertEqual([atom[0] for atom in structural[0].values()], ["NH4^+", "Cl^-"])
        displayed = draw.layout("[NH4+].[Cl-]", "displayed")
        displayed_labels = [atom[0] for atom in displayed[0].values()]
        self.assertEqual(displayed[0][0][0], "N^+")
        self.assertEqual(displayed[0][1][0], "Cl^-")
        self.assertEqual(set(displayed[0]) - {0, 1}, {2, 3, 4, 5})
        self.assertIn("N^+", displayed_labels)
        self.assertEqual(displayed_labels.count("H"), 4)
        self.assertIn("Cl^-", displayed_labels)
        markup, _, _ = svg.render_svg("[NH4+].[Cl-]")
        self.assertIn(">+</text>", markup)
        self.assertIn(">-</text>", markup)

    def test_invisible_skeletal_fragment_is_rejected_and_identified(self):
        cases = (("C.[Na+]", "Fragment 1 'C'"), ("[Na+].C", "Fragment 2 'C'"))
        for smiles, message in cases:
            with self.subTest(smiles=smiles):
                with self.assertRaisesRegex(draw.DisconnectedStructureError, message):
                    draw.layout(smiles, "skeletal")


class PublicEntryPointTests(unittest.TestCase):
    def test_structure_entry_points_share_the_typed_boundary(self):
        entry_points = (
            lambda: draw.layout("ethanoic acid"),
            lambda: draw.render("ethanoic acid"),
            lambda: svg.render_svg("ethanoic acid"),
            lambda: svg.render_svg_inner("ethanoic acid"),
        )
        for call in entry_points:
            with self.subTest(call=call):
                with self.assertRaises(draw.InvalidSmilesError):
                    call()

    def test_reaction_smiles_species_accept_ionic_fragments(self):
        svg_result = reaction.render_reaction_svg(["CC(=O)[O-].[Na+]"], ["CC(=O)O"])
        png_result = reaction.render_reaction_png(["CC(=O)[O-].[Na+]"], ["CC(=O)O"])
        self.assertIn(">+</text>", svg_result[0])
        self.assertGreater(svg_result[1], 0)
        self.assertGreater(png_result.width, 0)

    def test_reaction_rejects_unsupported_topology(self):
        with self.assertRaises(draw.UnsupportedTopologyError):
            reaction.render_reaction_svg(["c1ccc2ccccc2c1"], ["c1ccccc1"])

    def test_docx_structure_entry_point_uses_the_same_boundary(self):
        from docx import Document

        doc = Document()
        run = doc.add_paragraph().add_run()
        with self.assertRaises(draw.InvalidSmilesError):
            embed_docx.add_structure(doc, run, "ethanoic acid")

    def test_reaction_accepts_separate_species_and_explicit_literals(self):
        markup, width, height = reaction.render_reaction_svg(
            ["CCO", "$H2O"], ["CC=O", "$H2"], reagent="Cu", conditions="heat"
        )
        self.assertTrue(markup.startswith("<svg"))
        self.assertGreater(width, 0)
        self.assertGreater(height, 0)

    def test_structure_render_parses_smiles_once(self):
        parser = draw.Chem.MolFromSmiles
        with mock.patch.object(draw.Chem, "MolFromSmiles", wraps=parser) as wrapped:
            draw.render("CC(=O)O")
        self.assertEqual(wrapped.call_count, 1)

    def test_reaction_parses_each_structural_species_once(self):
        parser = draw.Chem.MolFromSmiles
        with mock.patch.object(draw.Chem, "MolFromSmiles", wraps=parser) as wrapped:
            reaction.render_reaction_svg(["CCO", "$2[O]"], ["CC=O", "$H2O"])
        self.assertEqual(wrapped.call_count, 2)


class SupportedOutputStabilityTests(unittest.TestCase):
    SINGLE_FRAGMENT_HASHES = {
        "C1CCCCC1": {
            "displayed": "1a5525551f14c859fd164ba59703ff3ea0549e6f95c52285598e1b6e7c379f1e",
            "structural": "1a5525551f14c859fd164ba59703ff3ea0549e6f95c52285598e1b6e7c379f1e",
            "skeletal": "1a5525551f14c859fd164ba59703ff3ea0549e6f95c52285598e1b6e7c379f1e",
        },
        "CC#N": {
            "displayed": "2f8726b528023e1fd8261d5a76fb0779e0b26d6c7f8009c6facfe8797d75b664",
            "structural": "cc81536b0d90a2f15e8f0a3bfc8dab6f8c571a53a62a428af6e40eec91444ac2",
            "skeletal": "99a70e64a029aa04b1535457a24f440cdabe1a8a47b52a8ef5bcc6b416376d8b",
        },
        "CC(=O)O": {
            "displayed": "45e009594425098df399744d85029e8831fd77632f11805c210bf872e40c162e",
            "structural": "cb1a893b29063515ccbc9ea82da9da69330be8ef10639ec1910422c5ad58c8ea",
            "skeletal": "2785881bc46683f9ce09b4596e1a4c21739591685aec2f2482d13bf25f1e580f",
        },
        "CC(=O)[O-]": {
            "displayed": "4779b292f206c51de53fa85b1ce4c533921dbdf3dec4e8c298cf9f5ac4a53df8",
            "structural": "1db26efd7d9c9942ded8f3709e7d498e15c4b323b57e1f349a9d5bc287c54678",
            "skeletal": "ecb383f9e95cd3c4e5cdeafb1b7e21dca49fea2315ab25a60060890931a86f13",
        },
        "CC(O)C": {
            "displayed": "9994eb57bfdcb8f5b6727a83f3fcd88a3a274e1d79b6c78b376b58ec2a1eae0e",
            "structural": "95486ca6cffa537d136a57271ab7fb98d55e3999e26a7d03a516233b3b408daf",
            "skeletal": "b772c81469a687cc79d965a90fd7e4dd7f418af2f2b51c5370f748fd691ea19e",
        },
        "CC(O)C(=O)O": {
            "displayed": "6e65a06bad9a46f674580f95ad6ff5b75761581838a3d496bdb42063fc30d7f3",
            "structural": "aed223b437b383376bd6143f239b7a62ea55fc348ccfc4a1b8eb3a10df261081",
            "skeletal": "264c53fe370ab332654b86c79c317d1ac685b1b07590c7bed2b0b4b8a0196391",
        },
        "CC=CC": {
            "displayed": "00caf4b2c76b5a5977a8365fe2dd667cd297431d97ecf4d2bd88cab834677d31",
            "structural": "0ef8c86c313a25a04a660e40d8a15d8a5a5850c5a5593b7a1d3e43a9db619065",
            "skeletal": "018e6b1ef9c26e4dd48c01350209e69e93fd0dde71f90fdb6ecbdc1e9e1c7249",
        },
        "O=[N+]([O-])c1ccccc1": {
            "displayed": "47a2c07a6700a622630bf1092a21cfa79dd6d45663d4e0dce6e90c60cb11aa62",
            "structural": "47a2c07a6700a622630bf1092a21cfa79dd6d45663d4e0dce6e90c60cb11aa62",
            "skeletal": "47a2c07a6700a622630bf1092a21cfa79dd6d45663d4e0dce6e90c60cb11aa62",
        },
        "[NH3+]CC(=O)[O-]": {
            "displayed": "22af0c097f2e465a1add6729822db72498fb34823fb81881075f43b001d27fc5",
            "structural": "c85a3cd303e61127e51ee7fdbc19b3d520d29bd2a42211965b503391f3dcef26",
            "skeletal": "272a9abb4b3bc815b788f1537588a936b294dacbf10e587a8391df72e1d8d5ce",
        },
    }
    STRUCTURES = {
        "ethanoic-acid": "CC(=O)O",
        "ethyl-ethanoate": "CCOC(C)=O",
        "nitrobenzene": "O=[N+]([O-])c1ccccc1",
    }
    REACTIONS = {
        "esterification": {
            "reactants": ["CCO", "CC(=O)O"],
            "products": ["CCOC(C)=O", "$H2O"],
            "reagent": "conc. H2SO4",
            "conditions": "heat",
            "arrow": "<=>",
        },
        "ethanol-oxidation": {
            "reactants": ["CCO", "$2[O]"],
            "products": ["CC(=O)O", "$H2O"],
            "reagent": "K2Cr2O7 / H2SO4",
            "conditions": "reflux",
            "arrow": "->",
        },
    }

    def test_manifest_single_fragment_renders_keep_pre_change_hashes(self):
        with open(os.path.join(ROOT, "evaluation", "preview_manifest.json")) as fh:
            manifest = json.load(fh)
        manifest_smiles = set()
        for case in manifest["cases"]:
            smiles = case["input"]
            if case["expected"]["outcome"] != "supported" or not isinstance(smiles, str):
                continue
            molecule = draw.Chem.MolFromSmiles(smiles)
            self.assertIsNotNone(molecule, case["id"])
            if len(draw.Chem.GetMolFrags(molecule)) == 1:
                manifest_smiles.add(smiles)
        self.assertEqual(set(self.SINGLE_FRAGMENT_HASHES), manifest_smiles)
        if draw.FONT_PATH != draw.ARIAL:
            self.skipTest("Exact pre-change hashes require the reference Arial font")
        for smiles, expected_by_form in self.SINGLE_FRAGMENT_HASHES.items():
            for form, expected in expected_by_form.items():
                with self.subTest(smiles=smiles, form=form):
                    actual = hashlib.sha256(draw.render(smiles, form=form).tobytes()).hexdigest()
                    self.assertEqual(actual, expected)

    def test_tracked_structure_examples_are_unchanged(self):
        for name, smiles in self.STRUCTURES.items():
            with self.subTest(name=name):
                markup, width, height = svg.render_svg(smiles)
                actual = draw.render(smiles)
                if draw.FONT_PATH == draw.ARIAL:
                    with open(os.path.join(ROOT, "examples", f"{name}.svg")) as fh:
                        self.assertEqual(markup, fh.read())
                    with Image.open(os.path.join(ROOT, "examples", f"{name}.png")) as tracked:
                        expected = tracked.convert("RGB")
                    self.assertEqual(actual.size, expected.size)
                    self.assertEqual(actual.tobytes(), expected.tobytes())
                else:
                    self._assert_portable_artifacts(markup, width, height, actual)

    def test_tracked_reaction_examples_are_unchanged(self):
        for name, kwargs in self.REACTIONS.items():
            with self.subTest(name=name):
                markup, width, height = reaction.render_reaction_svg(**kwargs)
                actual = reaction.render_reaction_png(**kwargs)
                if draw.FONT_PATH == draw.ARIAL:
                    with open(os.path.join(ROOT, "examples", f"{name}.svg")) as fh:
                        self.assertEqual(markup, fh.read())
                    with Image.open(os.path.join(ROOT, "examples", f"{name}.png")) as tracked:
                        expected = tracked.convert("RGB")
                    self.assertEqual(actual.size, expected.size)
                    self.assertEqual(actual.tobytes(), expected.tobytes())
                else:
                    self._assert_portable_artifacts(markup, width, height, actual)

    def _assert_portable_artifacts(self, markup, width, height, image):
        root = ET.fromstring(markup)
        self.assertTrue(root.tag.endswith("svg"))
        self.assertGreater(width, 0)
        self.assertGreater(height, 0)
        self.assertEqual(image.size, (width, height))
        white = Image.new(image.mode, image.size, "white")
        self.assertIsNotNone(ImageChops.difference(image, white).getbbox())


if __name__ == "__main__":
    unittest.main()
