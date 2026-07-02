"""Embed structure_draw structures into a .docx as sharp VECTOR (SVG) with a PNG fallback.

    from embed_docx import add_structure, build_doc
    build_doc([("propanone","CC(C)=O"), ("benzene","c1ccccc1")], "out.docx", title="My structures")

Needs: python-docx, rdkit, pillow (structure_draw + structure_svg in same folder).
"""
import os, sys, tempfile
sys.path.insert(0, os.path.dirname(__file__))
import structure_draw as N
import structure_svg as S
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.packuri import PackURI
from docx.opc.part import Part
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import parse_xml

_SVG_EXT_URI = "{96DAC541-7B7A-43D3-8B79-37D633B846F1}"
_counter = [0]

def add_structure(doc, run, smiles, width_in=1.7, png_path=None):
    """Render `smiles` and place it in `run` as SVG (vector) + PNG fallback."""
    if png_path is None:
        png_path = os.path.join(tempfile.gettempdir(), f"_mol_{_counter[0]}.png")
    N.render(smiles, png_path)                       # PNG fallback (also sets display size)
    svg, _, _ = S.render_svg(smiles)
    run.add_picture(png_path, width=Inches(width_in))
    part = Part(PackURI(f'/word/media/struct{_counter[0]}.svg'),
                'image/svg+xml', svg.encode('utf-8'), doc.part.package)
    _counter[0] += 1
    rid = doc.part.relate_to(part, RT.IMAGE)
    blip = run._r.xpath('.//a:blip')[-1]
    blip.append(parse_xml(
        '<a:extLst xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">'
        f'<a:ext uri="{_SVG_EXT_URI}">'
        '<asvg:svgBlip xmlns:asvg="http://schemas.microsoft.com/office/drawing/2016/SVG/main" '
        'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" '
        f'r:embed="{rid}"/></a:ext></a:extLst>'))

def build_doc(items, out_path, title=None, cols=4, width_in=1.7):
    """items: list of (caption, smiles) OR (group_title, [(caption, smiles), ...]).
    Writes a captioned grid of vector structures to out_path."""
    doc = Document()
    if title:
        doc.add_heading(title, 0)
    groups = items if items and isinstance(items[0][1], list) else [(None, items)]
    for gtitle, members in groups:
        if gtitle:
            doc.add_heading(gtitle, level=2)
        rows = (len(members) + cols - 1) // cols
        tbl = doc.add_table(rows=rows * 2, cols=cols)
        for idx, (cap, smi) in enumerate(members):
            r, c = (idx // cols) * 2, idx % cols
            p = tbl.cell(r, c).paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_structure(doc, p.add_run(), smi, width_in=width_in)
            pc = tbl.cell(r + 1, c).paragraphs[0]; pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = pc.add_run(cap); run.font.size = Pt(10)
    doc.save(out_path)
    return out_path
